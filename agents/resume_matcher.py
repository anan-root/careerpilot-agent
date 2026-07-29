"""Resume parsing, job matching, and targeted advice generation."""

from __future__ import annotations

import json
import re
import hashlib
from pathlib import Path
from typing import Iterable

from agents.prompt_loader import render_prompt
from llm_client import chat, chat_json

TEXT_EXTENSIONS = {".txt", ".md", ".tex"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}
_AI_MATCH_CACHE: dict[str, dict] = {}

STOPWORDS = {
    "and", "or", "the", "with", "for", "from", "this", "that", "you", "your",
    "岗位", "要求", "职责", "负责", "熟悉", "掌握", "优先", "相关", "能力",
    "工作", "项目", "经验", "公司", "职位", "实习", "开发", "工程师",
}

TECH_TERMS = [
    "Python", "Java", "JavaScript", "TypeScript", "C++", "Go", "SQL",
    "PyTorch", "TensorFlow", "Transformers", "HuggingFace", "LangChain",
    "LangGraph", "FastAPI", "Flask", "Django", "React", "Vue", "Docker",
    "Kubernetes", "Linux", "Git", "MySQL", "PostgreSQL", "Redis",
    "Elasticsearch", "MinIO", "Pandas", "NumPy", "Scikit-learn",
    "大模型", "LLM", "Agent", "AI Agent", "智能体", "RAG", "检索增强",
    "向量检索", "Embedding", "ReRanker", "BM25", "Prompt", "Function Calling",
    "MCP", "微调", "SFT", "DPO", "RLHF", "LoRA", "MoE", "Transformer",
    "Tokenizer", "多模态", "NLP", "CV", "推荐系统", "机器学习", "深度学习",
    "算法", "后端", "前端", "全栈", "爬虫", "数据分析", "数据挖掘",
    "实习", "校招", "工程化", "分布式", "高并发", "数据库", "API",
]


def extract_resume_text(file_path: str | Path) -> str:
    """Extract text from a resume file."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"暂不支持 {suffix}，请上传 PDF、DOCX、TXT、MD 或 TEX。")

    if suffix in TEXT_EXTENSIONS:
        return _normalize_text(path.read_text(encoding="utf-8", errors="ignore"))
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    raise ValueError(f"暂不支持 {suffix}。")


def analyze_resume(resume_text: str) -> dict:
    """Use the LLM to turn raw resume text into structured evidence."""
    prompt = render_prompt("resume_profile", resume_text=_clip(resume_text, 12000))
    try:
        return chat_json(prompt, system="你是严谨的招聘简历解析助手，输出必须是合法 JSON。")
    except Exception:
        return _fallback_resume_profile(resume_text)


def rank_jobs_for_resume(
    resume_text: str,
    jobs: Iterable[dict],
    *,
    top_n: int | None = 20,
    ai_top_n: int = 0,
    progress_callback=None,
    resume_cache_key: str | None = None,
) -> list[dict]:
    """Rank jobs by resume fit. Heuristic first, optional LLM refinement."""
    resume_keywords = set(_extract_keywords(resume_text))
    ranked: list[dict] = []

    for job in jobs:
        match = _score_job_against_resume(resume_text, resume_keywords, job)
        ranked.append({**job, "resume_match": match})

    ranked.sort(key=lambda j: j["resume_match"]["score"], reverse=True)

    if ai_top_n > 0:
        for index, job in enumerate(ranked[:ai_top_n], 1):
            cache_key = _ai_match_cache_key(resume_text, job, resume_cache_key=resume_cache_key)
            if cache_key in _AI_MATCH_CACHE:
                _emit_progress(
                    progress_callback,
                    f"DeepSeek 精排缓存命中第 {index}/{ai_top_n} 个：{job.get('company', '')} / {job.get('title', '')}",
                )
                ai_match = dict(_AI_MATCH_CACHE[cache_key])
            else:
                _emit_progress(
                    progress_callback,
                    f"DeepSeek 正在精排第 {index}/{ai_top_n} 个：{job.get('company', '')} / {job.get('title', '')}",
                )
                ai_match = deep_match_resume_to_job(resume_text, job)
                if isinstance(ai_match.get("score"), (int, float)):
                    _AI_MATCH_CACHE[cache_key] = dict(ai_match)
            job["resume_match"]["ai"] = ai_match
            job["ai_match"] = ai_match
            ai_score = ai_match.get("score")
            if isinstance(ai_score, (int, float)):
                job["resume_match"]["score"] = round((job["resume_match"]["score"] * 0.25) + (ai_score * 0.75), 1)
                job["resume_match"]["ai_refined"] = True
            else:
                job["resume_match"]["ai_refined"] = False
        ranked.sort(key=lambda j: j["resume_match"]["score"], reverse=True)
        _emit_progress(progress_callback, "DeepSeek 精排阶段结束")

    if top_n is not None:
        return ranked[:top_n]
    return ranked


def _emit_progress(callback, message: str) -> None:
    if callback is None:
        return
    try:
        callback(message)
    except Exception:
        pass


def _ai_match_cache_key(resume_text: str, job: dict, *, resume_cache_key: str | None = None) -> str:
    resume_hash = resume_cache_key or hashlib.sha1(_normalize_text(resume_text).encode("utf-8")).hexdigest()[:16]
    job_parts = [
        str(job.get(key) or "").strip()
        for key in ("platform", "job_id", "company", "title", "location", "salary")
    ]
    job_hash = hashlib.sha1("|".join(job_parts).encode("utf-8")).hexdigest()[:16]
    return f"{resume_hash}:{job_hash}"


def deep_match_resume_to_job(resume_text: str, job: dict) -> dict:
    """Ask the LLM for a strict resume-to-job fit score."""
    prompt = render_prompt(
        "job_match",
        resume_text=_clip(resume_text, 12000),
        job_json=json.dumps(_job_for_prompt(job), ensure_ascii=False, indent=2),
    )
    try:
        data = chat_json(prompt, system="你是严谨的招聘匹配评估助手，输出必须是合法 JSON。")
        return _normalize_ai_match(data)
    except Exception as exc:
        return {
            "score": None,
            "level": "AI分析失败",
            "matched_evidence": [],
            "missing_requirements": [],
            "missing_keywords": [],
            "risk_points": [str(exc)],
            "risks": [str(exc)],
            "resume_actions": [],
            "interview_focus": [],
            "reasoning": "",
        }


def generate_job_gap_analysis(resume_text: str, job: dict) -> str:
    """Generate a JD/resume gap analysis for one target job."""
    prompt = render_prompt(
        "job_gap_analysis",
        resume_text=_clip(resume_text, 14000),
        job_json=json.dumps(_job_for_prompt(job), ensure_ascii=False, indent=2),
    )
    try:
        return chat(prompt, system="你是专业但诚实的求职顾问，严格反对简历编造。", max_tokens=5000)
    except Exception as exc:
        return build_fallback_job_gap_analysis(resume_text, job, error=exc)


def generate_interview_pack(resume_text: str, job: dict) -> str:
    """Generate a structured interview preparation pack for one target job."""
    prompt = render_prompt(
        "interview_pack",
        resume_text=_clip(resume_text, 14000),
        job_json=json.dumps(_job_for_prompt(job), ensure_ascii=False, indent=2),
    )
    try:
        return chat(prompt, system="你是 AI 应用开发岗位面试教练，严格基于真实简历给建议。", max_tokens=6500)
    except Exception as exc:
        return build_fallback_interview_pack(resume_text, job, error=exc)


def generate_resume_job_advice(resume_text: str, job: dict) -> str:
    """Generate resume optimization and interview advice for one target job."""
    prompt = f"""{render_prompt(
        "job_gap_analysis",
        resume_text=_clip(resume_text, 14000),
        job_json=json.dumps(_job_for_prompt(job), ensure_ascii=False, indent=2),
    )}

另外请追加：
## 匹配结论
## 简历优化建议
## 可直接改写的简历要点
## 面试准备重点
## 面试官可能追问
## 投递前检查清单"""
    try:
        return chat(prompt, system="你是专业但诚实的求职顾问，严格反对简历编造。", max_tokens=6000)
    except Exception as exc:
        return build_fallback_resume_job_advice(resume_text, job, error=exc)


def build_fallback_resume_job_advice(resume_text: str, job: dict, error: Exception | None = None) -> str:
    """Generate deterministic advice when the LLM is unavailable."""
    resume_keywords = set(_extract_keywords(resume_text))
    job_text = _job_text(job)
    job_keywords = set(_extract_keywords(job_text))
    matched = sorted(resume_keywords & job_keywords, key=lambda item: (-len(item), item))[:12]
    missing = sorted(job_keywords - resume_keywords, key=lambda item: (-len(item), item))[:12]
    title = job.get("title", "")
    company = job.get("company", "")
    salary = job.get("salary", "")
    location = job.get("location", "")
    safe_error = error.__class__.__name__ if error else "LLMUnavailable"

    lines = [
        "# 简历优化意见和面试建议",
        "",
        "> DeepSeek 精评暂不可用，已自动切换为本地规则建议。请检查 API Key 或网络后可重新生成精评。",
        f"> 降级原因：{safe_error}",
        "",
        "## 匹配结论",
        "",
        f"- 目标岗位：{company} - {title}",
        f"- 地点/薪资：{location or '未知'} / {salary or '未知'}",
        f"- 简历已覆盖关键词：{', '.join(matched) or '暂未发现明显重合关键词'}",
        f"- 建议补强关键词：{', '.join(missing[:8]) or '暂无明显缺口'}",
        "",
        "## 简历优化建议",
        "",
        "- 把最相关项目放到简历前半部分，标题里直接出现岗位方向关键词。",
        "- 每条项目描述采用“场景-动作-结果”结构，优先写你真实做过的技术决策和产出。",
        "- 对岗位要求中已掌握的关键词，补充对应项目证据；没有做过的内容不要硬写。",
        "- 如果岗位偏 Agent/RAG，把检索、工具调用、上下文构建、服务封装、评估指标写清楚。",
        "",
        "## 可直接改写的简历要点",
        "",
    ]
    if matched:
        for keyword in matched[:5]:
            lines.append(f"- 围绕 `{keyword}` 补一条真实项目证据，例如你的职责、技术实现和结果。")
    else:
        lines.append("- 先补充 1 个与岗位方向最接近的项目证据，再投递该岗位。")

    lines.extend([
        "",
        "## 面试准备重点",
        "",
    ])
    focus_terms = [term for term in ("RAG", "Agent", "大模型", "LLM", "Prompt", "向量检索", "FastAPI", "React", "Python") if term.lower() in job_text.lower()]
    for term in focus_terms[:8] or ["项目架构", "技术选型", "问题定位", "结果量化"]:
        lines.append(f"- 准备 `{term}` 的项目实践、原理解释和常见追问。")

    lines.extend([
        "",
        "## 面试官可能追问",
        "",
        "- 这个项目为什么这样设计？你负责了哪一部分？",
        "- 遇到过什么线上或真实使用问题？你怎么定位和解决？",
        "- 如果数据量、并发量或准确率要求提升，你会怎么改？",
        "- 你简历中和岗位最相关的一段经历，能否用 2 分钟讲清楚？",
        "",
        "## 投递前检查清单",
        "",
        "- 简历里没有编造项目、学历、公司和结果。",
        "- 岗位关键词至少能在一个真实项目里找到证据。",
        "- 准备好 30 秒自我介绍和 2 分钟项目介绍。",
        "- 打开原岗位页面确认工作制、薪资、地点和公司真实性。",
    ])
    return "\n".join(lines)


def build_fallback_job_gap_analysis(resume_text: str, job: dict, error: Exception | None = None) -> str:
    """Generate deterministic JD/resume gap analysis when the LLM is unavailable."""
    resume_keywords = set(_extract_keywords(resume_text))
    job_keywords = set(_extract_keywords(_job_text(job)))
    matched = sorted(resume_keywords & job_keywords, key=lambda item: (-len(item), item))[:10]
    missing = sorted(job_keywords - resume_keywords, key=lambda item: (-len(item), item))[:10]
    safe_error = error.__class__.__name__ if error else "LLMUnavailable"

    lines = [
        "# JD/简历差距分析",
        "",
        f"> DeepSeek 暂不可用，已使用本地关键词规则生成。降级原因：{safe_error}",
        "",
        "## 岗位核心要求",
        "",
        f"- 岗位：{job.get('company', '')} - {job.get('title', '')}",
        f"- 关键词：{', '.join(sorted(job_keywords, key=lambda item: (-len(item), item))[:12]) or '暂无'}",
        "",
        "## 简历已覆盖内容",
        "",
        f"- 已命中：{', '.join(matched) or '暂未发现明显重合证据'}",
        "",
        "## 简历缺失内容",
        "",
        f"- 建议补证据：{', '.join(missing[:8]) or '暂无明显缺口'}",
        "",
        "## 可补充项目表达",
        "",
        "- 把最相关项目改成“业务场景-技术方案-个人动作-结果指标”的结构。",
        "- 对岗位关键词逐一补充真实证据，没有做过的内容标记为待学习，不要编造。",
        "",
        "## 技术栈补强建议",
        "",
        "- 优先补齐 Agent/RAG/LLM 应用开发的工程化表达：接口封装、检索链路、Prompt、评估和部署。",
        "",
        "## 投递风险提醒",
        "",
        "- 投递前打开原岗位页面确认薪资、经验、工作制和公司信息。",
    ]
    return "\n".join(lines)


def build_fallback_interview_pack(resume_text: str, job: dict, error: Exception | None = None) -> str:
    """Generate deterministic interview prep when the LLM is unavailable."""
    job_text = _job_text(job)
    focus_terms = [
        term for term in ("RAG", "Agent", "大模型", "LLM", "Prompt", "向量检索", "FastAPI", "Python", "React")
        if term.lower() in job_text.lower() or term.lower() in resume_text.lower()
    ]
    safe_error = error.__class__.__name__ if error else "LLMUnavailable"
    lines = [
        "# 面试准备包",
        "",
        f"> DeepSeek 暂不可用，已使用本地规则生成。降级原因：{safe_error}",
        "",
        "## 岗位理解",
        "",
        f"- 目标岗位：{job.get('company', '')} - {job.get('title', '')}",
        "- 面试重点通常会落在项目真实性、AI 应用落地、工程化能力和学习速度。",
        "",
        "## 简历追问预测",
        "",
        "- 你在最相关项目中具体负责什么？",
        "- 这个项目的难点是什么，如何验证效果？",
        "- 如果线上效果不好，你会怎么定位问题？",
        "",
        "## 技术面试题",
        "",
    ]
    for term in focus_terms[:8] or ["RAG", "Agent", "FastAPI", "Prompt"]:
        lines.append(f"- 请解释 `{term}` 在你项目中的作用、实现细节和可改进点。")
    lines.extend([
        "",
        "## 项目深挖问题",
        "",
        "- 项目架构为什么这样设计？",
        "- 数据、检索、模型调用、服务接口分别如何拆分？",
        "- 如何处理延迟、成本、稳定性和可观测性？",
        "",
        "## 行为面问题",
        "",
        "- 讲一次你快速学习并落地新技术的经历。",
        "- 讲一次项目受阻后你如何推进解决。",
        "",
        "## 反问面试官建议",
        "",
        "- 团队目前 AI 应用主要服务哪些业务场景？",
        "- 岗位更看重模型能力、后端工程还是产品落地？",
        "",
        "## 7 天准备清单",
        "",
        "- Day 1：重写最相关项目的 2 分钟介绍。",
        "- Day 2：梳理岗位 JD 和简历证据对应表。",
        "- Day 3：复盘 RAG/Agent/LLM 调用链路。",
        "- Day 4：准备项目难点、故障定位和优化方案。",
        "- Day 5：准备代码/接口/部署相关追问。",
        "- Day 6：模拟一轮技术面。",
        "- Day 7：确认岗位信息并准备反问问题。",
    ])
    return "\n".join(lines)


def build_match_report(ranked_jobs: list[dict]) -> str:
    """Create a local Markdown report for matched jobs."""
    lines = ["# 简历岗位匹配报告", ""]
    for i, job in enumerate(ranked_jobs, 1):
        match = job.get("resume_match", {})
        lines.extend([
            f"## {i}. {job.get('company', '')} - {job.get('title', '')}",
            f"- 匹配分：{match.get('score', 0)}",
            f"- 地点：{job.get('location', '')}",
            f"- 薪资：{job.get('salary', '')}",
            f"- 来源：{job.get('platform', '')}",
            f"- 命中关键词：{', '.join(match.get('matched_keywords', [])[:20]) or '暂无'}",
            f"- 缺口关键词：{', '.join(match.get('missing_keywords', [])[:20]) or '暂无'}",
            f"- 建议：{match.get('summary', '')}",
            "",
        ])
    return "\n".join(lines)


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("解析 PDF 需要安装 pypdf：pip install pypdf") from exc

    reader = PdfReader(str(path))
    parts = [page.extract_text() or "" for page in reader.pages]
    return _normalize_text("\n".join(parts))


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("解析 DOCX 需要安装 python-docx：pip install python-docx") from exc

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return _normalize_text("\n".join(parts))


def _fallback_resume_profile(text: str) -> dict:
    keywords = _extract_keywords(text)
    email = re.search(r"[\w.+-]+@[\w.-]+\.\w+", text)
    phone = re.search(r"(?:\+?86[- ]?)?1[3-9]\d{9}", text)
    return {
        "name": "",
        "target_role": "",
        "education": [],
        "skills": keywords[:40],
        "projects": [],
        "strengths": keywords[:15],
        "risks": ["未调用 AI 解析，仅使用本地关键词提取。"],
        "contact": {
            "email": email.group(0) if email else "",
            "phone": phone.group(0) if phone else "",
        },
    }


def _score_job_against_resume(resume_text: str, resume_keywords: set[str], job: dict) -> dict:
    job_text = _job_text(job)
    job_keywords = set(_extract_keywords(job_text))
    skill_keywords = set(_extract_keywords(" ".join(str(job.get(k, "")) for k in ("skills", "requirements", "description", "full_jd"))))

    matched = sorted(resume_keywords & job_keywords, key=lambda x: (-len(x), x))
    skill_matched = sorted(resume_keywords & skill_keywords, key=lambda x: (-len(x), x))
    missing = sorted(skill_keywords - resume_keywords, key=lambda x: (-len(x), x))

    keyword_score = _ratio_score(len(matched), max(len(job_keywords), 1), cap=0.45)
    skill_score = _ratio_score(len(skill_matched), max(len(skill_keywords), 1), cap=0.40)
    title_score = _title_score(resume_keywords, str(job.get("title", "")))
    description_bonus = 8 if any(k in resume_text for k in ("实习", "项目", "工程", "算法", "大模型", "Agent", "RAG")) else 0

    score = min(100.0, keyword_score + skill_score + title_score + description_bonus)
    quality_score = _job_quality_score(job)
    if quality_score is not None:
        if quality_score >= 75:
            score = min(100.0, score + 3)
        elif quality_score < 50:
            score = min(score, 68.0)

    if score >= 75:
        summary = "简历关键词和岗位要求重合较高，建议优先投递并做定制化优化。"
    elif score >= 55:
        summary = "有一定匹配基础，适合补强关键词和项目证据后投递。"
    elif score >= 35:
        summary = "匹配度一般，建议先判断是否真的符合目标方向。"
    else:
        summary = "当前简历证据较弱，不建议作为优先岗位。"

    return {
        "score": round(score, 1),
        "matched_keywords": matched[:30],
        "skill_matches": skill_matched[:20],
        "missing_keywords": missing[:30],
        "field_quality_score": quality_score,
        "summary": summary,
    }


def _ratio_score(matches: int, total: int, *, cap: float) -> float:
    if total <= 0:
        return 0.0
    return min(matches / total, cap) / cap * (cap * 100)


def _title_score(resume_keywords: set[str], title: str) -> float:
    title_keywords = set(_extract_keywords(title))
    if not title_keywords:
        return 0.0
    overlap = len(resume_keywords & title_keywords) / len(title_keywords)
    return min(overlap * 18, 18)


def _extract_keywords(text: str) -> list[str]:
    text = _normalize_text(text)
    raw = re.findall(r"[A-Za-z][A-Za-z0-9+#./-]{1,}", text)
    lower_text = text.lower()
    for term in TECH_TERMS:
        if term.lower() in lower_text:
            raw.append(term)
    for chunk in re.findall(r"[\u4e00-\u9fff]{2,12}", text):
        if len(chunk) <= 6:
            raw.append(chunk)
    items: list[str] = []
    seen: set[str] = set()
    for token in raw:
        item = token.strip().strip(".,;:()[]{}<>，。；：（）【】")
        key = item.lower()
        if len(item) < 2 or key in STOPWORDS:
            continue
        if key not in seen:
            seen.add(key)
            items.append(item)
    return items


def _job_text(job: dict) -> str:
    fields = [
        "company", "title", "location", "salary", "job_type", "skills",
        "degree", "experience", "description", "requirements", "full_jd",
        "company_industry", "company_stage", "welfare",
    ]
    return "\n".join(str(job.get(k, "")) for k in fields if job.get(k))


def _job_for_prompt(job: dict) -> dict:
    keys = [
        "company", "title", "location", "salary", "job_type", "skills",
        "degree", "experience", "description", "requirements", "full_jd",
        "platform", "url", "source_url", "company_size", "company_industry",
        "company_stage", "company_address", "welfare", "weekend_policy",
        "posted_date", "crawl_keyword",
    ]
    return {k: job.get(k, "") for k in keys if job.get(k)}


def _safe_score(value) -> float:
    try:
        return max(0.0, min(100.0, float(value)))
    except (TypeError, ValueError):
        return 0.0


def _job_quality_score(job: dict) -> float | None:
    value = job.get("field_quality_score")
    if value in (None, ""):
        return None
    try:
        return max(0.0, min(100.0, float(value)))
    except (TypeError, ValueError):
        return None


def _normalize_ai_match(data: dict) -> dict:
    score = _optional_score(data.get("score"))
    missing = _as_text_list(data.get("missing_requirements") or data.get("missing_keywords"))
    risks = _as_text_list(data.get("risk_points") or data.get("risks"))
    normalized = {
        "score": score,
        "level": _normalize_ai_level(str(data.get("level") or ""), score),
        "matched_evidence": _as_text_list(data.get("matched_evidence")),
        "missing_requirements": missing,
        "missing_keywords": missing,
        "risk_points": risks,
        "risks": risks,
        "resume_actions": _as_text_list(data.get("resume_actions")),
        "interview_focus": _as_text_list(data.get("interview_focus")),
        "reasoning": str(data.get("reasoning") or "").strip()[:220],
    }
    return normalized


def _normalize_ai_level(level: str, score: float | None) -> str:
    for item in ("强推", "推荐", "可投", "谨慎", "不建议"):
        if item in level:
            return item
    if score is None:
        return "AI分析失败"
    if score >= 90:
        return "强推"
    if score >= 80:
        return "推荐"
    if score >= 70:
        return "可投"
    if score >= 60:
        return "谨慎"
    if score >= 50:
        return "备选"
    return "不建议"


def _optional_score(value: object) -> float | None:
    if value in (None, ""):
        return None
    try:
        return max(0.0, min(100.0, float(value)))
    except (TypeError, ValueError):
        return None


def _as_text_list(value: object) -> list[str]:
    if value is None:
        return []
    if isinstance(value, (list, tuple, set)):
        source = value
    else:
        source = [value]
    result = []
    for item in source:
        text = str(item or "").strip()
        if text and text not in result:
            result.append(text)
    return result


def _normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n")).strip()


def _clip(text: str, limit: int) -> str:
    text = _normalize_text(text)
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[内容过长，已截断]"
